"""
Document processing service for PDF, DOCX, and TXT files.
"""

import os
from typing import List, Dict, Any, Optional
import hashlib


from app.config import get_settings
from app.models.database import Database
from app.core.vector_store import VectorStore
from app.services.text_splitter import LegislationTextSplitter


class DocumentService:
    """Service for document processing and management."""

    SUPPORTED_TYPES = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "text/plain": "txt",
    }

    def __init__(self, db: Database, vector_store: VectorStore):
        self.db = db
        self.vector_store = vector_store
        settings = get_settings()

        # Initialize legislation-aware text splitter
        self.text_splitter = LegislationTextSplitter(
            chunk_size=settings.chunk_size, chunk_overlap=settings.chunk_overlap
        )

    def calculate_hash(self, content: bytes) -> str:
        """Calculate MD5 hash of file content."""
        return hashlib.md5(content).hexdigest()

    async def process_file(
        self,
        file_content: bytes,
        filename: str,
        content_type: str,
        title: Optional[str] = None,
        description: Optional[str] = None,
        is_temporary: bool = False,
    ) -> Dict[str, Any]:
        """Process uploaded file and store embeddings."""
        # Validate file type
        if content_type not in self.SUPPORTED_TYPES:
            raise ValueError(f"Desteklenmeyen dosya tipi: {content_type}")

        # Calculate file hash
        file_hash = self.calculate_hash(file_content)

        # Check for existing document with same hash
        db = self.db.get_db()
        fs = self.db.get_fs()
        existing_doc = await db.documents.find_one({"file_hash": file_hash})
        if existing_doc:
            # If it exists, return it immediately
            return self.db._serialize_doc(existing_doc)

        file_type = self.SUPPORTED_TYPES[content_type]

        # Extract text from file
        text = await self._extract_text(file_content, file_type)

        if not text.strip():
            raise ValueError("Dosyadan metin çıkarılamadı")

        # Upload to GridFS
        # db.fs is AsyncIOMotorGridFSBucket
        grid_id = await fs.upload_from_stream(
            filename, file_content, metadata={"content_type": content_type}
        )

        # Create document record
        document = await self.db.create_document(
            {
                "title": title or filename,
                "description": description or "",
                "filename": filename,
                "file_type": file_type,
                "file_size": len(file_content),
                "gridfs_id": grid_id,  # Link to GridFS file
                "file_hash": file_hash,
                "is_temporary": is_temporary,
            }
        )

        try:
            # Update status to processing
            await self.db.update_document(document["id"], {"status": "processing"})

            # Split text into chunks with metadata
            chunks_data = self.text_splitter.split_text_with_metadata(text)
            chunks = [item["content"] for item in chunks_data]

            # Create metadata for each chunk
            metadata = []
            for i, item in enumerate(chunks_data):
                meta = {
                    "filename": filename,
                    "title": title or filename,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "document_id": document["id"],  # Explicitly add doc ID to filtering
                }
                if "metadata" in item:
                    meta.update(item["metadata"])
                metadata.append(meta)

            # Store chunks with embeddings
            chunk_count = await self.vector_store.add_documents(
                document_id=document["id"], chunks=chunks, metadata=metadata
            )

            # Update document with chunk count
            updated = await self.db.update_document(
                document["id"], {"chunk_count": chunk_count, "status": "completed"}
            )
            if updated is None:
                raise RuntimeError("Document update failed")

            return updated

        except Exception as e:
            # Update status to failed
            await self.db.update_document(
                document["id"], {"status": "failed", "error": str(e)}
            )
            # Cleanup GridFS file if processing failed
            await fs.delete(grid_id)
            raise

    async def _extract_text(self, content: bytes, file_type: str) -> str:
        """Extract text from file content."""
        if file_type == "txt":
            return self._extract_txt(content)
        elif file_type == "pdf":
            return self._extract_pdf(content)
        elif file_type == "docx":
            return self._extract_docx(content)
        else:
            raise ValueError(f"Desteklenmeyen dosya tipi: {file_type}")

    def _extract_txt(self, content: bytes) -> str:
        """Extract text from TXT file."""
        # Try different encodings
        encodings = ["utf-8", "utf-16", "windows-1254", "iso-8859-9", "latin-1"]

        for encoding in encodings:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue

        # Fallback with error handling
        return content.decode("utf-8", errors="ignore")

    def _extract_pdf(self, content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            from pypdf import PdfReader
            from io import BytesIO

            reader = PdfReader(BytesIO(content))
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            return "\n\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"PDF okunamadı: {str(e)}")

    def _extract_pdf_pages(self, content: bytes) -> list[dict]:
        """Extract text from PDF file page by page."""
        try:
            from pypdf import PdfReader
            from io import BytesIO

            reader = PdfReader(BytesIO(content))
            pages = []

            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                pages.append({"page": i + 1, "text": text})

            return pages
        except Exception as e:
            raise ValueError(f"PDF okunamadı: {str(e)}")

    def _extract_docx(self, content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            from io import BytesIO

            doc = Document(BytesIO(content))
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"DOCX okunamadı: {str(e)}")

    async def list_documents(
        self, skip: int = 0, limit: int = 50
    ) -> List[Dict[str, Any]]:
        """List all documents (excluding temporary ones)."""
        # We need to filter by is_temporary: False or Not Exists
        cursor = (
            self.db.get_db()
            .documents.find(
                {"$or": [{"is_temporary": False}, {"is_temporary": {"$exists": False}}]}
            )
            .sort("created_at", -1)
            .skip(skip)
            .limit(limit)
        )

        docs = await cursor.to_list(length=limit)
        return [self.db._serialize_doc(doc) for doc in docs]

    async def get_document(self, document_id: str) -> Optional[Dict[str, Any]]:
        """Get document by ID."""
        return await self.db.get_document(document_id)

    async def delete_document(self, document_id: str) -> bool:
        """Delete document and its embeddings/GridFS file."""
        # Get document to find GridFS ID
        doc = await self.db.get_document(document_id)

        # Delete from GridFS if linked
        if doc and "gridfs_id" in doc:
            try:
                await self.db.get_fs().delete(doc["gridfs_id"])
            except Exception:
                # Log error but continue deleting document record
                pass
        # Legacy cleanup: delete local file if exists (for old docs)
        elif doc and "file_path" in doc:
            try:
                os.remove(doc["file_path"])
            except OSError:
                pass

        return await self.db.delete_document(document_id)

    async def list_stored_files(self) -> List[Dict[str, Any]]:
        """List files in GridFS."""
        files = []
        cursor = self.db.get_fs().find().limit(50)
        async for grid_out in cursor:
            files.append(
                {
                    "name": grid_out.filename,
                    "size": grid_out.length,
                    "uploaded_at": grid_out.upload_date,
                    "id": str(grid_out._id),
                }
            )
        return files
